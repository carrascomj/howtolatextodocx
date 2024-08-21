import re
import bibtexparser
import docx
import typer
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

app = typer.Typer()


def parse_bibtex(bibfile: str) -> dict[str, str]:
    with open(bibfile, 'r') as bibtex_file:
        bib_database = bibtexparser.load(bibtex_file)
    return {entry['ID']: entry for entry in bib_database.entries}


def add_hyperlink(paragraph, url: str, text: str):
    """Add hyperlink to a `paragraph` with visible context `text` and `url`."""

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Set the hyperlink style (optional, depends on your template)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Join all the xml elements together and add the required text to the w:r element
    new_run.append(rPr)
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)

    # Append the hyperlink to the paragraph
    paragraph._p.append(hyperlink)


def format_author(entry: dict) -> str:
    """Format author names given a bib entry, depending on number of authors."""
    authors = entry.get('author', 'Unknown').split(',')
    author_format = authors[0]
    if len(authors) == 2:
        author_format = f"{authors[0]} and {authors[1]}"        
    if len(authors) > 2:
        author_format = f"{authors[0]} et al."
    return author_format


def replace_references(docfile: str, bib_entries: dict, outputfile: str):
    """Replace text in a docfile of the form @somebibref with a proper reference."""
    doc = docx.Document(docfile)
    pattern = re.compile(r'\[\-?@([^\]]+)\]')

    for paragraph in doc.paragraphs:
        matches = pattern.findall(paragraph.text)
        if matches:
            new_text = paragraph.text
            citations = []
            urls = []
            for match in matches:
                # split multiple citations within the same block
                ref_ids = match.split('; ')
                formatted_citations = []
                i = 0
                for ref_id in ref_ids:
                    if i != 0:
                        # remove starting @
                        ref_id = ref_id[1:]
                    i += 1
                    if ref_id in bib_entries:
                        author = format_author(bib_entries[ref_id])
                        year = bib_entries[ref_id].get('year', 'n.d.')
                        citation = f"{author} {year}"
                        formatted_citations.append(citation)
                        doi = bib_entries[ref_id].get('doi', '')
                        url_bib = bib_entries[ref_id].get('url', '')
                        url = f"https://doi.org/{doi}" if doi else url_bib  if url_bib else ""
                        urls.append(url)

                # join formatted citations and replace in the text
                citations += formatted_citations
                formatted_citation_text = '; '.join(formatted_citations)
                new_text = new_text.replace(f"[@{match}]", f" ({formatted_citation_text})")
                new_text = new_text.replace(f"[-@{match}]", formatted_citation_text)

            # add paragraph run by run (piece of text, citation, piece of text, etc)
            paragraph.clear()
            rest = new_text
            for cit, url in zip(citations, urls):
                splits = rest.split(cit, maxsplit=1)
                if len(splits) > 1:
                    text, rest = splits
                    paragraph.add_run(text)
                    if url:
                        add_hyperlink(paragraph, url, cit)
                    else:
                        paragraph.add_run(cit)
                else:
                    rest = splits[0]
            paragraph.add_run(rest)

    doc.save(outputfile)


@app.command()
def replace(
    bib_file: str = typer.Argument(..., help="Path to the input .bib file"),
    doc_file: str = typer.Argument(..., help="Path to the input .docx file"),
    output_file: str = typer.Argument(..., help="Path to the output .docx file")
):
    """Replace references from a docx generated from a md generated from tex."""
    bib_entries = parse_bibtex(bib_file)
    replace_references(doc_file, bib_entries, output_file)
    typer.echo(f"References replaced and saved in {output_file}")


if __name__ == "__main__":
    app()
